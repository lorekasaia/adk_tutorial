import os
import uuid
import re
import pandas as pd
import matplotlib.pyplot as plt
from google import adk
from database import consultar_cloud_sql, MAPA_ESTADOS

def obtener_resumen_pipeline() -> str:
    try:
        df = consultar_cloud_sql("")
        if df.empty:
            return "No hay datos suficientes para generar un resumen."
        
        if '_estado' in df.columns and 'valor_estimado' in df.columns:
            df['estado_texto'] = df['_estado'].map(MAPA_ESTADOS).fillna('Desconocido')
            resumen = df.groupby('estado_texto').agg(
                cantidad_clientes=('id', 'count'),
                valor_total_mxn=('valor_estimado', 'sum'),
                ticket_promedio=('valor_estimado', 'mean')
            ).reset_index()
            
            return "Resumen Financiero del Pipeline:\n" + resumen.to_json(orient="records", force_ascii=False)
        else:
            return "La base de datos no contiene las columnas de valor o estado necesarias para este análisis."
    except Exception as e:
        return f"Error al generar el resumen estadístico: {e}"

def consultar_dashboard_bi(kpi: str, contexto: str = "general") -> str:
    print(f"[BI API Mock] Solicitando KPI: '{kpi}' con contexto '{contexto}'")
    if kpi == "ventas_totales":
        return f"El Dashboard de BI reporta que las ventas totales para '{contexto}' son de $1,450,000 MXN este trimestre."
    elif kpi == "tasa_conversion":
        return f"Según la plataforma de BI, la tasa de conversión de prospectos a clientes para '{contexto}' se sitúa en un 24.5%."
    elif kpi == "rendimiento_vendedores":
        return "El reporte de BI indica que Lore lidera las ventas con un 120% de alcance de cuota, seguida de cerca por el resto del equipo."
    else:
        return f"Datos del Dashboard para el KPI '{kpi}': Los indicadores están estables y dentro de los rangos esperados para el período actual."

def generar_grafico_analisis(metrica: str) -> str:
    """
    Genera un gráfico visual basado en los datos de clientes.
    El parámetro 'metrica' debe ser una de las siguientes opciones (o un sinónimo cercano):
    - 'estado': Gráfico circular de la distribución de clientes por estado en el pipeline.
    - 'prioridad': Gráfico de barras de clientes por nivel de prioridad.
    - 'valor': Gráfico de barras del valor estimado del pipeline por estado.
    - 'fuente': Gráfico de barras horizontales del origen de los prospectos.
    - 'conversion': Gráfico circular de la tasa de conversión general.
    """
    try:
        df = consultar_cloud_sql("") 
        if df.empty:
            return "No hay suficientes datos en la base para graficar."
        
        # --- Normalización de métrica para flexibilidad ---
        # Permite que el usuario pida "estado actual" y se mapee a "estado".
        metrica_limpia = metrica.lower()
        if "estado" in metrica_limpia:
            metrica_limpia = "estado"
        elif "prioridad" in metrica_limpia:
            metrica_limpia = "prioridad"
        elif "valor" in metrica_limpia:
            metrica_limpia = "valor"
        elif "fuente" in metrica_limpia or "origen" in metrica_limpia:
            metrica_limpia = "fuente"
        elif "conversión" in metrica_limpia or "conversion" in metrica_limpia:
            metrica_limpia = "conversion"
        # --- Fin de la normalización ---

        plt.figure(figsize=(8, 6))
        if metrica_limpia == "prioridad" and 'prioridad' in df.columns:
            conteo = df['prioridad'].value_counts()
            conteo.plot(kind='bar', color=['#4CAF50', '#FF9800', '#F44336'])
            plt.title("Distribución de Clientes por Prioridad")
            plt.xlabel("Nivel de Prioridad")
            plt.ylabel("Cantidad de Clientes")
            plt.xticks(rotation=0)
        elif metrica_limpia == "estado" and '_estado' in df.columns:
            estados_texto = df['_estado'].map(MAPA_ESTADOS).fillna('Otro (' + df['_estado'].astype(str) + ')')
            conteo = estados_texto.value_counts()
            conteo.plot(kind='pie', autopct='%1.1f%%', startangle=90)
            plt.title("Proporción de Clientes por Estado Interno")
            plt.ylabel("")
        elif metrica_limpia == "valor" and 'valor_estimado' in df.columns and '_estado' in df.columns:
            df['estado_texto'] = df['_estado'].map(MAPA_ESTADOS).fillna('Otro')
            suma_valor = df.groupby('estado_texto')['valor_estimado'].sum().sort_values(ascending=False)
            suma_valor.plot(kind='bar', color='#2196F3')
            plt.title("Valor Estimado ($) del Pipeline por Estado")
            plt.xlabel("Estado del Cliente")
            plt.ylabel("Valor Estimado Total")
            plt.xticks(rotation=45, ha='right')
        elif metrica_limpia == "fuente" and 'fuente' in df.columns:
            conteo = df['fuente'].fillna('Desconocido').value_counts()
            conteo.sort_values().plot(kind='barh', color='#9C27B0')
            plt.title("Origen de los Prospectos (Fuentes)")
            plt.xlabel("Cantidad de Clientes")
            plt.ylabel("Fuente")
        elif metrica_limpia == "conversion" and 'es_cliente' in df.columns:
            conteo = df['es_cliente'].map({True: 'Cliente Cerrado', False: 'Prospecto Activo'}).value_counts()
            conteo.plot(kind='pie', autopct='%1.1f%%', startangle=90, colors=['#00BCD4', '#FFC107'])
            plt.title("Tasa de Conversión General")
            plt.ylabel("")
        else:
            return f"No se pudo generar el gráfico para la métrica '{metrica}'. Las métricas permitidas son: 'prioridad', 'estado', 'valor', 'fuente' y 'conversion'. Si el usuario pide 'todas', debes ejecutar esta herramienta 5 veces seguidas (una por cada métrica permitida)."
        
        filename = f"grafico_{uuid.uuid4().hex[:8]}.png"
        filepath = os.path.join("graficos", filename)
        plt.savefig(filepath, bbox_inches='tight')
        plt.close()
        return f"Gráfico generado con éxito. DEBES responder esto al usuario exactamente así para que vea la imagen: <br><img src='/graficos/{filename}' alt='Gráfico de {metrica}' style='max-width: 100%; border-radius: 8px; margin-top: 10px;'/>"
    except Exception as e:
        return f"Error al procesar y graficar los datos: {e}"

def exportar_datos_excel(termino_busqueda: str = "") -> str:
    try:
        df = consultar_cloud_sql(termino_busqueda)
        if df.empty:
            return "No hay datos en la base de datos que coincidan con ese criterio para exportar."
        
        filename = f"reporte_clientes_{uuid.uuid4().hex[:8]}.xlsx"
        filepath = os.path.join("reportes", filename)
        df.to_excel(filepath, index=False)
        return f"Reporte Excel generado con éxito. DEBES responder esto al usuario para que pueda descargarlo: <br><a href='/reportes/{filename}' download style='display: inline-block; padding: 10px 15px; background: #107c41; color: white; text-decoration: none; border-radius: 5px; margin-top: 10px; font-weight: bold;'>📊 Descargar Reporte Excel</a>"
    except Exception as e:
        return f"Error al generar el archivo Excel: {e}"

def generar_reporte_pdf(titulo: str, contenido: str, nombre_imagen: str = "") -> str:
    """
    Crea un documento PDF con un título, un contenido de texto y opcionalmente una imagen.
    Si generaste un gráfico antes, pasa su nombre de archivo (ej. 'grafico_123.png') en 'nombre_imagen'.
    Guarda el archivo en la carpeta 'reportes' y devuelve un enlace de descarga.
    """
    try:
        from fpdf import FPDF
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", 'B', size=16)
        pdf.cell(0, 10, txt=titulo, ln=True, align='C')
        pdf.ln(10)

        if nombre_imagen:
            match = re.search(r'grafico_[a-zA-Z0-9]+\.png', nombre_imagen)
            if match:
                ruta_img = os.path.join("graficos", match.group(0))
                if os.path.exists(ruta_img):
                    pdf.image(ruta_img, w=160)
                    pdf.ln(10)

        pdf.set_font("Arial", size=12)
        # El texto debe estar codificado para evitar errores con caracteres especiales en FPDF
        contenido_encoded = contenido.encode('latin-1', 'replace').decode('latin-1')
        pdf.multi_cell(0, 10, txt=contenido_encoded)
        
        filename = f"reporte_{uuid.uuid4().hex[:8]}.pdf"
        filepath = os.path.join("reportes", filename)
        pdf.output(filepath)
        
        return f"Reporte PDF generado. Responde con este enlace para descarga: <br><a href='/reportes/{filename}' download style='display: inline-block; padding: 10px 15px; background: #D32F2F; color: white; text-decoration: none; border-radius: 5px; margin-top: 10px; font-weight: bold;'>📄 Descargar Reporte PDF</a>"
    except Exception as e:
        return f"Error al generar el archivo PDF: {e}"

def generar_reporte_word(titulo: str, contenido: str, nombre_imagen: str = "") -> str:
    """
    Crea un documento Word (.docx) con un título, un contenido de texto y opcionalmente una imagen.
    Si generaste un gráfico antes, pasa su nombre de archivo (ej. 'grafico_123.png') en 'nombre_imagen'.
    Guarda el archivo en la carpeta 'reportes' y devuelve un enlace de descarga.
    """
    try:
        from docx import Document
        from docx.shared import Inches
        document = Document()
        document.add_heading(titulo, level=1)

        if nombre_imagen:
            match = re.search(r'grafico_[a-zA-Z0-9]+\.png', nombre_imagen)
            if match:
                ruta_img = os.path.join("graficos", match.group(0))
                if os.path.exists(ruta_img):
                    document.add_picture(ruta_img, width=Inches(6.0))

        document.add_paragraph(contenido)
        
        filename = f"reporte_{uuid.uuid4().hex[:8]}.docx"
        filepath = os.path.join("reportes", filename)
        document.save(filepath)
        
        return f"Reporte Word generado. Responde con este enlace para descarga: <br><a href='/reportes/{filename}' download style='display: inline-block; padding: 10px 15px; background: #2B579A; color: white; text-decoration: none; border-radius: 5px; margin-top: 10px; font-weight: bold;'>📄 Descargar Reporte Word</a>"
    except Exception as e:
        return f"Error al generar el archivo Word: {e}"

analytics_agent = adk.Agent(
    name="AnalyticsAgent",
    model="gemini-2.5-flash",
    instruction="Eres un analista de datos y BI. Tu propósito es generar resúmenes financieros, crear gráficos (<img>), exportar a Excel (<a>), consultar KPIs y generar reportes en formato PDF o Word. IMPORTANTE: Si el usuario te pide incluir un gráfico en un reporte PDF o Word, PRIMERO debes ejecutar la herramienta 'generar_grafico_analisis', leer el nombre del archivo generado en su respuesta (ej. grafico_xxx.png), y LUEGO ejecutar la herramienta de reporte pasándole ese nombre en el parámetro 'nombre_imagen'.",
    tools=[generar_grafico_analisis, obtener_resumen_pipeline, exportar_datos_excel, consultar_dashboard_bi, generar_reporte_pdf, generar_reporte_word]
)